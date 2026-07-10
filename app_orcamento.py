import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
import json
import os
import io
import gspread
from google.oauth2.service_account import Credentials
import re

# =============================================================================
# 🧱 1. CONFIGURAÇÕES INICIAIS E FUNÇÕES UTILITÁRIAS DE SEGURANÇA
# =============================================================================
st.set_page_config(page_title="Sistema Lenoor - Orçamentos", page_icon="⚙️", layout="wide")

ID_PLANILHA = "1QrYGJCY-NwsRnMgOe1eugQmssu3mYsIc7wMxBT2oBSk"

DEFAULTS_MAQUINAS = {
    "CNC": 150.0, "Torno Automático": 120.0, "Freza": 120.0,
    "Furadeira": 70.0, "Torno Revolver": 90.0, "Retífica": 150.0,
    "Serra": 70.0, "Montagem": 90.0, "Outro": 150.0
}
DEFAULTS_IMPOSTOS = {"IR": 6.0, "FS": 4.0}

DEFAULTS_MATERIAIS = {
    "aço sx": {"constante": 0.68, "preco_atual": 0.0, "data_cotacao": ""}, 
    "aço red": {"constante": 0.62, "preco_atual": 0.0, "data_cotacao": ""}, 
    "aço quad": {"constante": 0.72, "preco_atual": 0.0, "data_cotacao": ""},
    "aluminio red": {"constante": 0.212, "preco_atual": 0.0, "data_cotacao": ""}, 
    "aluminio quad": {"constante": 0.212, "preco_atual": 0.0, "data_cotacao": ""}, 
    "aluminio sx": {"constante": 0.212, "preco_atual": 0.0, "data_cotacao": ""},
    "latao red": {"constante": 0.68, "preco_atual": 0.0, "data_cotacao": ""}, 
    "latao quad": {"constante": 0.78, "preco_atual": 0.0, "data_cotacao": ""}, 
    "latao sx": {"constante": 0.72, "preco_atual": 0.0, "data_cotacao": ""}
}

COLUNAS_PADRAO = [
    "Data/Hora", "Origem/Alteração", "Código da Peça", "Nome da Peça",
    "Empresa", "Comprador", "Lote", "Comprimento (mm)", "Margem Corte (mm)",
    "Tipo MP", "Preço MP Unitário", "Liga", "Diâmetro (mm)", "Diâmetro Externo (mm)",
    "Diâmetro Interno (mm)", "Total Kg Lote", "Custo Material (R$)", "Custo Tratamento (R$)",
    "Margem Lucro (%)", "Preço Total Lote (R$)", "Preço Unitário (R$)", "Usinagem_JSON", "Tratamento_JSON"
]

COLUNAS_COTACOES = [
    "ID Cotação", "Data", "Cliente", "Código", "Nome", "Lote", 
    "Preço de Venda Unitário", "Total Faturado", "Valor Total Proposta"
]

def safe_float(val, default=0.0):
    if pd.isna(val) or val is None or str(val).strip() == "":
        return default
    if isinstance(val, (int, float)):
        return float(val)
    
    # Limpador agressivo: Remove R$, %, espaços
    val_str = str(val).upper().replace("R$", "").replace("%", "").strip()
    val_str = val_str.replace(" ", "")
    
    if not val_str:
        return default
        
    # Trata conversão regional (ex: 1.500,50 -> 1500.50 ou 30,00 -> 30.00)
    if '.' in val_str and ',' in val_str:
        val_str = val_str.replace('.', '').replace(',', '.')
    elif ',' in val_str:
        val_str = val_str.replace(',', '.')
        
    try: 
        return float(val_str)
    except: 
        return default

def safe_int(val, default=0):
    if pd.isna(val) or val is None or str(val).strip() == "":
        return default
    try: 
        return int(float(val))
    except: 
        return default

def safe_str(val, default=""):
    return default if pd.isna(val) or val is None else str(val).strip()

def conectar_sheets():
    scope = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
    creds_dict = st.secrets["gcp_service_account"]
    creds = Credentials.from_service_account_info(creds_dict, scopes=scope)
    client = gspread.authorize(creds)
    return client

def ler_aba_sheets(nome_aba, colunas_padrao):
    try:
        client = conectar_sheets()
        sh = client.open_by_key(ID_PLANILHA)
        worksheet = sh.worksheet(nome_aba)
        dados = worksheet.get_all_records(value_render_option='UNFORMATTED_VALUE')
        if not dados:
            return pd.DataFrame(columns=colunas_padrao)
        df = pd.DataFrame(dados)
        for col in colunas_padrao:
            if col not in df.columns: 
                df[col] = None
        return df
    except Exception as e:
        st.error(f"⚠️ Erro ao ler a aba '{nome_aba}' no Google Sheets: {e}")
        return pd.DataFrame(columns=colunas_padrao)

def salvar_aba_sheets(df, nome_aba):
    try:
        client = conectar_sheets()
        sh = client.open_by_key(ID_PLANILHA)
        worksheet = sh.worksheet(nome_aba)
        worksheet.clear()
        lista_dados = [df.columns.values.tolist()] + df.fillna("").values.tolist()
        worksheet.update(range_name="A1", values=lista_dados, value_input_option="USER_ENTERED")
        return True
    except Exception as e:
        st.error(f"❌ Erro ao salvar dados na aba '{nome_aba}' do Google Sheets: {e}")
        return False

def carregar_config_nuvem():
    try:
        client = conectar_sheets()
        sh = client.open_by_key(ID_PLANILHA)
        
        try:
            wks_maq = sh.worksheet("Config_Maquinas")
            dados_maq = wks_maq.get_all_records(value_render_option='UNFORMATTED_VALUE')
        except: dados_maq = []
            
        if not dados_maq:
            rows = []
            for k, v in DEFAULTS_MAQUINAS.items(): rows.append({"Tipo": "Máquina", "Nome": k, "Valor": float(v)})
            for k, v in DEFAULTS_IMPOSTOS.items(): rows.append({"Tipo": "Imposto", "Nome": k, "Valor": float(v)})
            df_m = pd.DataFrame(rows)
            salvar_aba_sheets(df_m, "Config_Maquinas")
            valores_maquinas = DEFAULTS_MAQUINAS.copy()
            impostos = DEFAULTS_IMPOSTOS.copy()
        else:
            valores_maquinas, impostos = {}, {}
            for r in dados_maq:
                tipo = safe_str(r.get("Tipo"))
                nome = safe_str(r.get("Nome"))
                valor = safe_float(r.get("Valor"))
                if "Imposto" in tipo: impostos[nome] = valor
                else: valores_maquinas[nome] = valor

        try:
            wks_mat = sh.worksheet("Config_Materiais")
            dados_mat = wks_mat.get_all_records(value_render_option='UNFORMATTED_VALUE')
        except: dados_mat = []
            
        if not dados_mat:
            rows = []
            for k, v in DEFAULTS_MATERIAIS.items():
                rows.append({"Liga": k, "Constante": float(v["constante"]), "Preço Atual (R$/Kg)": float(v["preco_atual"]), "Data da Cotação": ""})
            df_mt = pd.DataFrame(rows)
            salvar_aba_sheets(df_mt, "Config_Materiais")
            materiais = DEFAULTS_MATERIAIS.copy()
        else:
            materials = {}
            for r in dados_mat:
                liga = safe_str(r.get("Liga"))
                constante_val = safe_float(r.get("Constante", r.get("constante")))
                preco_val = safe_float(r.get("Preço Atual (R$/Kg)", r.get("preco_atual")))
                data_val = safe_str(r.get("Data da Cotação", r.get("data_cotacao")))
                materials[liga] = {"constante": constante_val, "preco_atual": preco_val, "data_cotacao": data_val}
            materiais = materials
                
        return {"valores_maquinas": valores_maquinas, "impostos": impostos, "materiais": materiais}
    except Exception as e:
        st.error(f"Erro crítico no handshake com o Google Sheets: {e}")
        return {"valores_maquinas": DEFAULTS_MAQUINAS, "impostos": DEFAULTS_IMPOSTOS, "materiais": DEFAULTS_MATERIAIS}

# =============================================================================
# 🧠 2. INICIALIZAÇÃO BLINDADA DO COFRE (SESSION STATE - MEMÓRIA LOCAL CONTRA 429)
# =============================================================================
# Inicialização primária de Widgets para extinguir em definitivo os KeyErrors
valores_padrao_widgets = {
    "sel_empresa_aba1": "", "txt_comprador": "", "sel_codigo_peca": "", "txt_nome_peca": "",
    "txt_novo_codigo": "", "txt_nova_empresa": "",
    "num_lote": 100, "num_comprimento": 0.0, "num_margem_corte": 5.0, 
    "sel_tipo_mp": "Por Peso (Barra Maciça/Sextavada)", "num_preco_mp": 0.0, "sel_liga": "aço sx",
    "num_diam_barra": 15.0, "num_di_ext": 20.0, "num_di_int": 10.0, "num_peso_metro": 10.0,
    "num_peso_peca": 10.0, "num_peso_fornecido": 0.0, "slider_lucro_aba1": 30
}
for chave, val in valores_padrao_widgets.items():
    if chave not in st.session_state: st.session_state[chave] = val

if "menu_anterior" not in st.session_state: st.session_state.menu_anterior = ""
if "editor_version" not in st.session_state: st.session_state.editor_version = 0  
if "taxas_original_msg" not in st.session_state: st.session_state["taxas_original_msg"] = ""
if "msg_sucesso_aba1" not in st.session_state: st.session_state["msg_sucesso_aba1"] = ""
if "msg_sucesso_aba3" not in st.session_state: st.session_state["msg_sucesso_aba3"] = ""
if "msg_sucesso_aba4" not in st.session_state: st.session_state["msg_sucesso_aba4"] = ""

if "df_usinagem_v3" not in st.session_state:
    st.session_state["df_usinagem_v3"] = pd.DataFrame([{"Operação": "Tornear", "Máquina": "Torno Automático", "Peças por Hora": 50.0}])
if "df_tratamento_v3" not in st.session_state:
    st.session_state["df_tratamento_v3"] = pd.DataFrame([{"Tratamento": "Nenhum / Sem Tratamento", "Preço por Kg (R$)": 0.0}])

# CARREGAMENTO EM CACHE LOCAL: Só lê as tabelas do Google uma única vez!
if "db_historico" not in st.session_state or "db_cotacoes" not in st.session_state:
    st.session_state["db_historico"] = ler_aba_sheets("Historico", COLUNAS_PADRAO)
    st.session_state["db_cotacoes"] = ler_aba_sheets("Historico_Cotacoes", COLUNAS_COTACOES)
    config_carregada = carregar_config_nuvem()
    st.session_state.valores_maquinas = config_carregada["valores_maquinas"]
    st.session_state.impostos = config_carregada["impostos"]
    st.session_state.materiais = config_carregada["materiais"]

def limpar_formulario_orcamento():
    for k, v in valores_padrao_widgets.items(): st.session_state[k] = v
    st.session_state["df_usinagem_v3"] = pd.DataFrame([{"Operação": "Tornear", "Máquina": "Torno Automático", "Peças por Hora": 50.0}])
    st.session_state["df_tratamento_v3"] = pd.DataFrame([{"Tratamento": "Nenhum / Sem Tratamento", "Preço por Kg (R$)": 0.0}])
    st.session_state["taxas_original_msg"] = ""
    st.session_state.editor_version += 1

# Referencia as variáveis locais em memória RAM (Adeus lentidão e erro 429)
df_init = st.session_state["db_historico"]
lista_empresas = sorted(df_init["Empresa"].dropna().astype(str).unique().tolist()) if not df_init.empty else []
empresa_atual = st.session_state.get("sel_empresa_aba1", "")
if empresa_atual and empresa_atual != "➕ Novo Cliente...":
    lista_codigos = sorted(df_init[df_init["Empresa"] == empresa_atual]["Código da Peça"].dropna().astype(str).unique().tolist())
else:
    lista_codigos = sorted(df_init["Código da Peça"].dropna().astype(str).unique().tolist()) if not df_init.empty else []

def carregar_roteiro_antigo_callback():
    codigo_target = st.session_state.get("sel_codigo_peca", "")
    if codigo_target and codigo_target != "➕ Novo Código...":
        try:
            df_hist = st.session_state["db_historico"]
            df_filtrado = df_hist[df_hist["Código da Peça"] == codigo_target].copy()
            if not df_filtrado.empty:
                df_filtrado['_dt_temp'] = pd.to_datetime(df_filtrado["Data/Hora"], format="%d/%m/%Y %H:%M", errors='coerce')
                last_hist = df_filtrado.sort_values('_dt_temp').iloc[-1].fillna("").to_dict()
                
                st.session_state["sel_empresa_aba1"] = safe_str(last_hist.get("Empresa"))
                st.session_state["txt_comprador"] = safe_str(last_hist.get("Comprador"))
                st.session_state["txt_nome_peca"] = safe_str(last_hist.get("Nome da Peça"))
                st.session_state["num_lote"] = max(1, safe_int(last_hist.get("Lote"), default=100))
                st.session_state["num_comprimento"] = max(0.0, safe_float(last_hist.get("Comprimento (mm)"), default=0.0))
                st.session_state["num_margem_corte"] = max(0.0, safe_float(last_hist.get("Margem Corte (mm)"), default=5.0))
                st.session_state["sel_tipo_mp"] = safe_str(last_hist.get("Tipo MP", "Por Peso (Barra Maciça/Sextavada)"))
                
                liga_salva = safe_str(last_hist.get("Liga", "aço sx"))
                st.session_state["sel_liga"] = liga_salva
                preco_hist = safe_float(last_hist.get("Preço MP Unitário"), default=0.0)
                st.session_state["num_preco_mp"] = st.session_state.materiais.get(liga_salva, {}).get("preco_atual", preco_hist) if "Por Peso" in st.session_state["sel_tipo_mp"] else preco_hist

                st.session_state["num_diam_barra"] = max(0.0, safe_float(last_hist.get("Diâmetro (mm)"), default=15.0))
                st.session_state["num_di_ext"] = max(0.0, safe_float(last_hist.get("Diâmetro Externo (mm)"), default=20.0))
                st.session_state["num_di_int"] = max(0.0, safe_float(last_hist.get("Diâmetro Interno (mm)"), default=10.0))
                st.session_state["slider_lucro_aba1"] = max(15, min(safe_int(last_hist.get("Margem Lucro (%)"), default=30), 95))
                
                if "Usinagem_JSON" in last_hist and safe_str(last_hist["Usinagem_JSON"]) not in ["", "[]"]:
                    df_usi = pd.read_json(io.StringIO(str(last_hist["Usinagem_JSON"])), orient='records')
                    for col in ["Operação", "Máquina", "Peças por Hora"]:
                        if col not in df_usi.columns: df_usi[col] = "Outro" if col == "Máquina" else (50.0 if col == "Peças por Hora" else "")
                    
                    if "Preço/Hora_Aplicado" in df_usi.columns:
                        resumo = " | ".join([f"{r['Máquina']}: R$ {safe_float(r['Preço/Hora_Aplicado']):.2f}/h" for _, r in df_usi.drop_duplicates(subset=['Máquina']).iterrows()])
                        st.session_state["taxas_original_msg"] = f"Últimas taxas de máquina aplicadas no orçamento desta peça: {resumo}"
                    else: st.session_state["taxas_original_msg"] = ""
                    st.session_state["df_usinagem_v3"] = df_usi[["Operação", "Máquina", "Peças por Hora"]].copy()
                
                if "Tratamento_JSON" in last_hist and safe_str(last_hist["Tratamento_JSON"]) not in ["", "[]"]:
                    df_trat = pd.read_json(io.StringIO(str(last_hist["Tratamento_JSON"])))
                    st.session_state["df_tratamento_v3"] = df_trat[["Tratamento", "Preço por Kg (R$)"]].copy()
                    
                st.session_state.editor_version += 1
        except Exception as e: st.toast(f"Erro no autocomplete: {e}", icon="⚠️")

def atualizar_preco_mp():
    liga = st.session_state.get("sel_liga")
    if liga in st.session_state.materiais:
        st.session_state["num_preco_mp"] = safe_float(st.session_state.materiais[liga].get("preco_atual", 0.0))

# Atribuição local rápida das tabelas de custos
valores_maquinas = st.session_state.valores_maquinas
impostos = st.session_state.impostos
materiais = st.session_state.materiais

# =============================================================================
# 🧭 3. NAVEGAÇÃO LATERAL E TÍTULO
# =============================================================================
st.sidebar.title("🧭 Menu Lenoor")
opcoes = ["📊 1. Novo Orçamento", "📜 2. Histórico de Peças", "🧱 3. Matéria-Prima", "⚙️ 4. Custos Fixos & BD", "🛒 5. Cotação Final", "📁 6. Histórico de Cotações"]
opcao_menu = st.sidebar.radio("Navegação:", opcoes)

# Botão mestre de atualização manual para caso altere algo externo no Sheets celular
if st.sidebar.button("🔄 Forçar Sincronização Google Sheets", use_container_width=True):
    st.session_state["db_historico"] = ler_aba_sheets("Historico", COLUNAS_PADRAO)
    st.session_state["db_cotacoes"] = ler_aba_sheets("Historico_Cotacoes", COLUNAS_COTACOES)
    config_carregada = carregar_config_nuvem()
    st.session_state.valores_maquinas = config_carregada["valores_maquinas"]
    st.session_state.impostos = config_carregada["impostos"]
    st.session_state.materiais = config_carregada["materiais"]
    st.toast("⚡ Planilhas sincronizadas com sucesso!", icon="📥")
    st.rerun()

if opcao_menu != st.session_state.menu_anterior:
    st.session_state["taxas_original_msg"] = ""
    st.session_state["msg_sucesso_aba1"] = ""
    st.session_state["msg_sucesso_aba3"] = ""
    st.session_state["msg_sucesso_aba4"] = ""
    st.session_state.menu_anterior = opcao_menu

titulo_limpo = opcao_menu.split('. ')[1] if '. ' in opcao_menu else opcao_menu
st.title(titulo_limpo)
st.markdown("---")

# =============================================================================
# 📊 TELA 1: NOVO ORÇAMENTO
# =============================================================================
if opcao_menu == "📊 1. Novo Orçamento":
    st.subheader("👤 Dados do Cliente")
    col_cli1, col_cli2 = st.columns(2)
    with col_cli1:
        opces_empresa = [""] + lista_empresas + ["➕ Novo Cliente..."]
        if st.session_state["sel_empresa_aba1"] not in opces_empresa:
            opces_empresa.insert(1, st.session_state["sel_empresa_aba1"])
        empresa_sel = st.selectbox("Empresa/Cliente", opces_empresa, key="sel_empresa_aba1")
        
        if empresa_sel == "➕ Novo Cliente...":
            empresa = st.text_input("Novo Cliente", placeholder="Ex: TS", key="txt_nova_empresa").strip()
        else:
            empresa = empresa_sel
            
    with col_cli2:
        comprador = st.text_input("Comprador", key="txt_comprador", placeholder="Ex: Guilherme")

    st.markdown("---")
    st.subheader("📦 Dados do Produto")
    col_cod_linha1, col_cod_linha2 = st.columns([3, 1])
    with col_cod_linha1:
        opcoes_codigo = [""] + lista_codigos + ["➕ Novo Código..."]
        if st.session_state["sel_codigo_peca"] not in opcoes_codigo:
            opcoes_codigo.insert(1, st.session_state["sel_codigo_peca"])
        codigo_sel = st.selectbox("Código da Peça", opcoes_codigo, key="sel_codigo_peca")
        
        if codigo_sel == "➕ Novo Código...":
            codigo_peca = st.text_input("Novo Código:", placeholder="Ex: CH-0002", key="txt_novo_codigo").strip()
        else:
            codigo_peca = codigo_sel

    with col_cod_linha2:
        st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True) 
        if codigo_sel and codigo_sel != "➕ Novo Código...":
            st.button("Completar Dados", type="primary", use_container_width=True, on_click=carregar_roteiro_antigo_callback)

    if st.session_state["taxas_original_msg"]: st.info(st.session_state["taxas_original_msg"])

    col_dados_p1, col_dados_p2, col_dados_p3, col_dados_p4 = st.columns(4)
    with col_dados_p1:
        nome_peca = st.text_input("Nome da Peça", key="txt_nome_peca")
    with col_dados_p2:
        lote = st.number_input("Lote", min_value=1, step=1, key="num_lote")
    with col_dados_p3:
        comprimento = st.number_input("Comprimento (mm)", min_value=0.0, step=0.1, key="num_comprimento")
    with col_dados_p4:
        margem_corte = st.number_input("Corte/Perda (mm)", min_value=0.0, step=0.1, key="num_margem_corte")

    st.markdown("---")
    st.subheader("🧱 Matéria-Prima")
    opcoes_mp = ["Por Peso (Barra Maciça/Sextavada)", "Por Peso (Tubo/Bucha)", "Por Metro Linear", "Por Peça Pronta", "Mão de Obra Pura (Fornecido)"]
    tipo_mp = st.selectbox("Contabilização", opcoes_mp, key="sel_tipo_mp")
    
    custo_total_material, total_quilos, preco_unitario_mp, material_sel = 0.0, 0.0, 0.0, "N/A"
    diametro, di_ext, di_int = 0.0, 0.0, 0.0

    if tipo_mp == "Mão de Obra Pura (Fornecido)":
        total_quilos = st.number_input("Peso total lote (kg) [Tratamento]", min_value=0.0, step=0.1, key="num_peso_fornecido")
    else:
        col_mat1, col_mat2, col_mat3 = st.columns(3)

        if "Por Peso" in tipo_mp:
            with col_mat1:
                opcoes_liga = list(materiais.keys())
                material_sel = st.selectbox("Liga", opcoes_liga, key="sel_liga", on_change=atualizar_preco_mp)
                constante = materiais.get(material_sel, {}).get("constante", 0.0)
            
            with col_mat2:
                preco_unitario_mp = st.number_input("Preço Atual MP (R$)", min_value=0.0, step=1.0, key="num_preco_mp")
            
            if "Maciça" in tipo_mp:
                with col_mat3:
                    diametro = st.number_input("Diâmetro da Barra (mm)", min_value=0.0, step=0.1, key="num_diam_barra")
                peso_por_metro = ((diametro ** 2) * constante) / 100
            else:
                with col_mat3:
                    di_ext = st.number_input("Diâm. Ext. (mm)", min_value=0.0, step=0.1, key="num_di_ext")
                    di_int = st.number_input("Diâm. Int. (mm)", min_value=0.0, step=0.1, key="num_di_int")
                peso_por_metro = max(0.0, ((di_ext ** 2) - (di_int ** 2)) * constante / 100)

            total_metros = ((comprimento + margem_corte) * lote) / 1000
            total_quilos = total_metros * peso_por_metro
            custo_total_material = total_quilos * preco_unitario_mp

        elif tipo_mp == "Por Metro Linear":
            with col_mat1:
                preco_unitario_mp = st.number_input("Preço Atual MP (R$)", min_value=0.0, step=1.0, key="num_preco_mp")
            total_metros = ((comprimento + margem_corte) * lote) / 1000
            custo_total_material = total_metros * preco_unitario_mp
            with col_mat2:
                total_quilos = st.number_input("Peso total lote (kg)", min_value=0.0, step=0.1, key="num_peso_metro")

        elif tipo_mp == "Por Peça Pronta":
            with col_mat1:
                preco_unitario_mp = st.number_input("Preço Atual MP (R$)", min_value=0.0, step=1.0, key="num_preco_mp")
            custo_total_material = lote * preco_unitario_mp
            with col_mat2:
                total_quilos = st.number_input("Peso total lote (kg)", min_value=0.0, step=0.1, key="num_peso_peca")

    st.metric("Custo Total da Matéria-Prima", f"R$ {custo_total_material:.2f}", help=f"{total_quilos:.2f} kg")
    st.markdown("---")

    # --- PROCESSOS DE USINAGEM ---
    st.subheader("🔨 Usinagem")
    df_usinagem_input = st.data_editor(
        st.session_state["df_usinagem_v3"],
        num_rows="dynamic",
        column_config={
            "Operação": st.column_config.TextColumn("Operação", required=True),
            "Máquina": st.column_config.SelectboxColumn("Máquina", options=list(valores_maquinas.keys()), required=True),
            "Peças por Hora": st.column_config.NumberColumn("Produção (Pç/h)", min_value=0.1, step=1.0, default=50.0)
        },
        use_container_width=True,
        key=f"editor_usi_{st.session_state.editor_version}"
    )
    st.session_state["df_usinagem_v3"] = df_usinagem_input
    
    custo_total_usinagem = 0.0
    if not df_usinagem_input.empty:
        df_usi_clean = df_usinagem_input.fillna({"Peças por Hora": 50.0, "Máquina": "Outro"})
        for _, row in df_usi_clean.iterrows():
            pcs_h = max(0.1, safe_float(row.get("Peças por Hora"), 50.0))
            custo_total_usinagem += (lote / pcs_h) * valores_maquinas.get(safe_str(row.get("Máquina"), "Outro"), 120.0)

    st.caption(f"Custo de Usinagem: **R$ {custo_total_usinagem:.2f}**")
    st.markdown("---")

    # --- TRATAMENTOS ---
    st.subheader("✨ Tratamentos (Preço/KG)")
    df_trat_input = st.data_editor(
        st.session_state["df_tratamento_v3"],
        num_rows="dynamic",
        column_config={
            "Tratamento": st.column_config.TextColumn("Tratamento", required=True),
            "Preço por Kg (R$)": st.column_config.NumberColumn("Valor por Kg", min_value=0.0, step=0.01, format="R$ %.2f", default=0.0)
        },
        use_container_width=True,
        key=f"editor_trat_{st.session_state.editor_version}"
    )
    st.session_state["df_tratamento_v3"] = df_trat_input

    soma_preco_kg_tratamento = 0.0
    if not df_trat_input.empty:
        df_trat_clean = df_trat_input.fillna({"Preço por Kg (R$)": 0.0})
        for _, row in df_trat_clean.iterrows():
            soma_preco_kg_tratamento += safe_float(row.get("Preço por Kg (R$)"))
            
    custo_total_tratamentos = soma_preco_kg_tratamento * total_quilos
    st.caption(f"Custo Tratamento: **R$ {custo_total_tratamentos:.2f}**")
    st.markdown("---")

    # --- FECHAMENTO E SALVAMENTO ---
    porcentagem_lucro = st.slider("Margem de Lucro (%)", 15, 95, int(st.session_state["slider_lucro_aba1"]), 5)
    st.session_state["slider_lucro_aba1"] = porcentagem_lucro

    custo_bruto = custo_total_material + custo_total_usinagem + custo_total_tratamentos
    fator_lucro = max(0.05, (1 - (porcentagem_lucro / 100)))
    valor_venda_bruto = custo_bruto / fator_lucro
    total_imposto_pct = safe_float(impostos.get("IR")) + safe_float(impostos.get("FS"))
    valor_impostos = valor_venda_bruto * (total_imposto_pct / 100)
    preco_lote = valor_venda_bruto + valor_impostos

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Custo de Fábrica", f"R$ {custo_bruto:.2f}")
    c2.metric("Lucro Bruto", f"R$ {valor_venda_bruto - custo_bruto:.2f}")
    c3.metric(f"Impostos ({total_imposto_pct}%)", f"R$ {valor_impostos:.2f}")
    c4.metric("Preço UNITÁRIO Final", f"R$ {preco_lote / lote if lote > 0 else 0.0:.2f}")

    if st.button("💾 Gravar no Histórico", type="primary"):
        if not codigo_peca or codigo_peca == "➕ Novo Código...":
            st.error("Preencha o código da peça!")
        else:
            df_usi_salvar = df_usinagem_input.copy()
            if not df_usi_salvar.empty: df_usi_salvar["Preço/Hora_Aplicado"] = df_usi_salvar["Máquina"].map(valores_maquinas)
            
            novo_registro = {
                "Data/Hora": datetime.now().strftime("%d/%m/%Y %H:%M"), "Origem/Alteração": "Novo Orçamento",
                "Código da Peça": codigo_peca, "Nome da Peça": nome_peca, "Empresa": empresa, "Comprador": comprador,
                "Lote": lote, "Comprimento (mm)": comprimento, "Margem Corte (mm)": margem_corte, "Tipo MP": tipo_mp,
                "Preço MP Unitário": preco_unitario_mp, "Liga": material_sel, "Diâmetro (mm)": diametro,
                "Diâmetro Externo (mm)": di_ext, "Diâmetro Interno (mm)": di_int, "Total Kg Lote": round(total_quilos, 3),
                "Custo Material (R$)": round(custo_total_material, 2), "Custo Tratamento (R$)": round(custo_total_tratamentos, 2),
                "Margem Lucro (%)": porcentagem_lucro, "Preço Total Lote (R$)": round(preco_lote, 2),
                "Preço Unitário (R$)": round(preco_lote / lote if lote > 0 else 0.0, 2),
                "Usinagem_JSON": df_usi_salvar.to_json(orient='records'),
                "Tratamento_JSON": df_trat_input.to_json(orient='records')
            }
            df_final = pd.concat([df_init, pd.DataFrame([novo_registro])], ignore_index=True) if not df_init.empty else pd.DataFrame([novo_registro])
            
            sucesso = salvar_aba_sheets(df_final, "Historico")
            if sucesso:
                st.session_state["db_historico"] = df_final  # Atualiza cache local instantaneamente
                limpar_formulario_orcamento()
                st.session_state["msg_sucesso_aba1"] = f"✅ Orçamento da peça '{codigo_peca}' gravado com sucesso!"
                st.rerun()

    # CORREÇÃO DE FLUXO VISUAL: O aviso verde agora brota exatamente embaixo do botão de clique
    if st.session_state.get("msg_sucesso_aba1"):
        st.success(st.session_state["msg_sucesso_aba1"])
        st.session_state["msg_sucesso_aba1"] = ""

    st.markdown("---")
    with st.expander("❓ Entenda os Cálculos do Sistema (Fórmulas e Regras)"):
        st.markdown("#### 🧱 1. Cálculo de Tubos e Buchas")
        st.latex(r"Peso/Metro = \frac{(D_{ext}^2 - D_{int}^2) \times Constante\_da\_Liga}{100}")
        st.markdown("#### 🔨 2. Custo de Usinagem")
        st.latex(r"Custo\_Usinagem = \left( \frac{Tamanho\_do\_Lote}{Pe\text{\c{c}}as\_por\_Hora} \right) \times Taxa\_M\acute{a}quina\_(R\$/h)")
        st.markdown("#### 💰 3. Formação do Preço de Venda (Markup Divisor)")
        st.latex(r"Faturamento\_Bruto = \frac{Custo\_de\_F\acute{a}brica}{1 - \left( \frac{Margem\%}{100} \right)}")

# =============================================================================
# 📜 TELA 2: HISTÓRICO DE PEÇAS COM FILTRO EM CASCATA
# =============================================================================
elif opcao_menu == "📜 2. Histórico de Peças":
    if df_init.empty: 
        st.info("Nenhum orçamento gerado.")
    else:
        c_filt1, c_filt2 = st.columns(2)
        with c_filt1: cliente_filtro = st.multiselect("🔍 Cliente:", options=lista_empresas)
        
        todos_os_codigos = sorted(df_init["Código da Peça"].dropna().astype(str).unique())
        codigos_disp = sorted(df_init[df_init["Empresa"].isin(cliente_filtro)]["Código da Peça"].dropna().astype(str).unique()) if cliente_filtro else todos_os_codigos
        with c_filt2: codigo_filtro = st.multiselect("🔍 Código:", options=codigos_disp)
        
        opcao_visao = st.radio("Visualização:", ["🎯 Preços Atuais (Última Versão)", "⏳ Histórico Completo"])
        
        df_ord = df_init.copy()
        df_ord['_dt'] = pd.to_datetime(df_ord["Data/Hora"], format="%d/%m/%Y %H:%M", errors='coerce')
        df_ord = df_ord.sort_values("_dt", ascending=True)
        if cliente_filtro: df_ord = df_ord[df_ord["Empresa"].isin(cliente_filtro)]
        if codigo_filtro: df_ord = df_ord[df_ord["Código da Peça"].isin(codigo_filtro)]

        cols = ["Data/Hora", "Origem/Alteração", "Empresa", "Código da Peça", "Nome da Peça", "Lote", "Preço Unitário (R$)", "Preço Total Lote (R$)"]
        for c in cols: 
            if c not in df_ord.columns: df_ord[c] = "N/A"
        
        if "Completo" in opcao_visao: 
            df_exibir_t2 = df_ord[cols].iloc[::-1]
        else:
            df_ult = df_ord.groupby("Código da Peça").last().reset_index().sort_values("Empresa")
            df_exibir_t2 = df_ult[["Empresa", "Código da Peça", "Nome da Peça", "Lote", "Preço Unitário (R$)", "Preço Total Lote (R$)"]]
            
        # SOLUÇÃO DO PRINT: O configurador de colunas coloca a vírgula de forma visual mantendo o float puro no Python!
        st.dataframe(
            df_exibir_t2, 
            use_container_width=True,
            column_config={
                "Preço Unitário (R$)": st.column_config.NumberColumn("Preço Unit. (R$)", format="R$ %.2f"),
                "Preço Total Lote (R$)": st.column_config.NumberColumn("Preço Total (R$)", format="R$ %.2f")
            }
        )
        
        buffer_t2 = io.BytesIO()
        with pd.ExcelWriter(buffer_t2, engine='xlsxwriter') as writer:
            df_exibir_t2.to_excel(writer, sheet_name='Historico_Pecas', index=False)
        st.download_button("📥 Exportar Tabela para Excel (.xlsx)", data=buffer_t2.getvalue(), file_name="historico_pecas_lenoor.xlsx", mime="application/vnd.ms-excel")

# =============================================================================
# 🧱 TELA 3: MATÉRIA-PRIMA E RECÁLCULO INTELIGENTE
# =============================================================================
elif opcao_menu == "🧱 3. Matéria-Prima":
    st.write("Atualize os preços. A data de cotação atualiza automaticamente.")
    
    df_mat = pd.DataFrame.from_dict(materiais, orient="index").reset_index()
    df_mat.rename(columns={"index": "Liga", "preco_atual": "Preço Atual (R$/Kg)", "data_cotacao": "Data da Cotação"}, inplace=True)
    
    df_mat_edit = st.data_editor(
        df_mat,
        column_config={
            "Liga": st.column_config.TextColumn(disabled=True),
            "constante": st.column_config.NumberColumn("Constante Peso", disabled=True),
            "Preço Atual (R$/Kg)": st.column_config.NumberColumn(min_value=0.0, step=0.1, format="R$ %.2f"),
            "Data da Cotação": st.column_config.TextColumn(disabled=True)
        },
        use_container_width=True, hide_index=True
    )

    if st.button("💾 Salvar Novos Preços de Materiais", type="primary"):
        novos_materiais = {}
        mudancas = False
        hoje = datetime.now().strftime("%d/%m/%Y")
        ARQUIVO_AUDITORIA_MP = "historico_compras_mp.csv"
        registros_auditoria = []

        for _, row in df_mat_edit.iterrows():
            liga = row["Liga"]
            preco_novo = float(row["Preço Atual (R$/Kg)"])
            preco_antigo = materiais[liga]["preco_atual"]
            
            if preco_novo != preco_antigo: 
                mudancas = True
                registros_auditoria.append({"Data Alteração": hoje, "Liga/Material": liga, "Preço Antigo (R$)": preco_antigo, "Preço Novo (R$)": preco_novo})
                
            novos_materiais[liga] = {
                "constante": float(row["constante"]), 
                "preco_atual": preco_novo, 
                "data_cotacao": hoje if preco_novo != preco_antigo else materiais[liga]["data_cotacao"]
            }
            
        if mudancas:
            st.session_state.materiais = novos_materiais
            
            rows_mat = []
            for k, v in novos_materiais.items():
                rows_mat.append({"Liga": k, "Constante": float(v["constante"]), "Preço Atual (R$/Kg)": float(v["preco_atual"]), "Data da Cotação": safe_str(v["data_cotacao"])})
            
            sucesso_mat = salvar_aba_sheets(pd.DataFrame(rows_mat), "Config_Materiais")
            if sucesso_mat:
                if registros_auditoria:
                    df_auditoria_novo = pd.DataFrame(registros_auditoria)
                    if os.path.exists(ARQUIVO_AUDITORIA_MP):
                        try:
                            df_auditoria_antigo = pd.read_csv(ARQUIVO_AUDITORIA_MP)
                            df_auditoria_final = pd.concat([df_auditoria_antigo, df_auditoria_novo], ignore_index=True)
                        except: df_auditoria_final = df_auditoria_novo
                    else: df_auditoria_final = df_auditoria_novo
                    df_auditoria_final.to_csv(ARQUIVO_AUDITORIA_MP, index=False)
                    
                st.toast("✅ Configurações de materiais atualizadas na Nuvem!")
                st.rerun()

    ARQUIVO_AUDITORIA_MP = "historico_compras_mp.csv"
    with st.expander("📈 Ver Histórico de Evolução de Preços (MP)"):
        if not os.path.exists(ARQUIVO_AUDITORIA_MP):
            st.info("Nenhum histórico de alteração de preços registrado ainda.")
        else:
            try:
                df_aud = pd.read_csv(ARQUIVO_AUDITORIA_MP)
                materiais_historico = sorted(df_aud["Liga/Material"].unique().tolist())
                mat_grafico = st.selectbox("Escolha o material para ver o gráfico de evolução:", materiais_historico, key="recalc_graf_mp")
                df_filtrado_graf = df_aud[df_aud["Liga/Material"] == mat_grafico].copy()
                
                if not df_filtrado_graf.empty:
                    st.markdown(f"#### Tendência de Preço - {mat_grafico}")
                    df_graf_exibir = df_filtrado_graf.set_index("Data Alteração")[["Preço Novo (R$)"]]
                    st.line_chart(df_graf_exibir, use_container_width=True)
                    
                    st.markdown("#### Lista de Alterações")
                    st.dataframe(df_filtrado_graf.iloc[::-1], hide_index=True, use_container_width=True)
                    
                    buffer_mp = io.BytesIO()
                    with pd.ExcelWriter(buffer_mp, engine='xlsxwriter') as writer:
                        df_filtrado_graf.to_excel(writer, sheet_name='Historico_MP', index=False)
                    st.download_button("📥 Baixar Histórico (.xlsx)", data=buffer_mp.getvalue(), file_name=f"historico_precos_{mat_grafico}.xlsx", mime="application/vnd.ms-excel")
            except: st.error("Erro ao carregar o arquivo de histórico de compras de MP.")
            
    st.markdown("---")
    st.subheader("🔄 Recalcular por Material")
    if not df_init.empty:
        liga_recalc = st.selectbox("Selecione a Liga que sofreu alteração:", [""] + list(materiais.keys()))
        if liga_recalc:
            df_ultimos = df_init.sort_values(by="Data/Hora").groupby("Código da Peça").last().reset_index()
            df_afetados = df_ultimos[(df_ultimos["Liga"] == liga_recalc) & (df_ultimos["Tipo MP"].str.contains("Peso"))].copy()
            
            if df_afetados.empty: st.info("Nenhuma peça mapeada por peso usa esta liga.")
            else:
                st.warning("⚠️ O recálculo utilizará os valores SALVOS no sistema em nuvem.")
                st.write(f"Preço de referência para **{liga_recalc}**: R$ {materiais[liga_recalc]['preco_atual']:.2f}")
                df_afetados.insert(0, "Recalcular", True)
                
                df_afetados_edit = st.data_editor(
                    df_afetados[["Recalcular", "Empresa", "Código da Peça", "Nome da Peça", "Lote", "Total Kg Lote"]],
                    hide_index=True, use_container_width=True,
                    column_config={
                        "Recalcular": st.column_config.CheckboxColumn("Recalcular?"),
                        "Empresa": st.column_config.TextColumn(disabled=True),
                        "Código da Peça": st.column_config.TextColumn(disabled=True),
                        "Nome da Peça": st.column_config.TextColumn(disabled=True),
                        "Lote": st.column_config.NumberColumn(disabled=True),
                        "Total Kg Lote": st.column_config.NumberColumn(disabled=True)
                    }
                )
                
                if st.button("🚀 Executar Recálculo de MP"):
                    codigos_selecionados = df_afetados_edit[df_afetados_edit["Recalcular"]]["Código da Peça"].tolist()
                    novas_linhas = []
                    novo_preco_mp = materiais[liga_recalc]['preco_atual']
                    
                    for _, row in df_afetados[df_afetados["Código da Peça"].isin(codigos_selecionados)].iterrows():
                        row_dict = row.to_dict()
                        row_dict["Preço MP Unitário"] = novo_preco_mp
                        row_dict["Custo Material (R$)"] = round(safe_float(row_dict.get("Total Kg Lote", 0.0)) * novo_preco_mp, 2)
                        
                        roteiro_raw = row_dict.get("Usinagem_JSON", "[]")
                        roteiro = json.loads(roteiro_raw) if isinstance(roteiro_raw, str) else []
                        custo_usinagem_antigo = 0.0
                        lote_val = max(1, safe_int(row_dict.get("Lote", 100)))
                        
                        if isinstance(roteiro, list):
                            for op in roteiro:
                                pcs_h = max(0.1, safe_float(op.get("Peças por Hora"), 50.0))
                                maq_nome = safe_str(op.get("Máquina"), "Outro")
                                taxa_hora = safe_float(op.get("Preço/Hora_Aplicado", valores_maquinas.get(maq_nome, 120.0)))
                                custo_usinagem_antigo += (lote_val / pcs_h) * taxa_hora
                        
                        custo_bruto = row_dict["Custo Material (R$)"] + custo_usinagem_antigo + safe_float(row_dict.get("Custo Tratamento (R$)"))
                        fator_l = max(0.05, (1 - (safe_float(row_dict.get("Margem Lucro (%)", 30)) / 100)))
                        valor_venda = custo_bruto / fator_l
                        imp_pct = safe_float(impostos.get("IR")) + safe_float(impostos.get("FS"))
                        preco_final = valor_venda + (valor_venda * (imp_pct / 100))
                        
                        row_dict["Preço Total Lote (R$)"] = round(preco_final, 2)
                        row_dict["Preço Unitário (R$)"] = round(preco_final / lote_val, 2)
                        row_dict["Origem/Alteração"] = "Recálculo de MP"
                        row_dict["Data/Hora"] = datetime.now().strftime("%d/%m/%Y %H:%M")
                        
                        if "Recalcular" in row_dict: del row_dict["Recalcular"]
                        if "_dt_temp" in row_dict: del row_dict["_dt_temp"]
                        novas_linhas.append(row_dict)
                        
                    if novas_linhas:
                        df_final_salvar = pd.concat([df_init, pd.DataFrame(novas_linhas)], ignore_index=True)
                        sucesso_recalc = salvar_aba_sheets(df_final_salvar, "Historico")
                        if sucesso_recalc:
                            st.session_state["db_historico"] = df_final_salvar
                            st.session_state["msg_sucesso_aba3"] = f"✅ {len(novas_linhas)} peças recalculadas com sucesso pela troca de MP!"
                            st.rerun()

    if st.session_state.get("msg_sucesso_aba3"):
        st.success(st.session_state["msg_sucesso_aba3"])
        st.session_state["msg_sucesso_aba3"] = ""

# =============================================================================
# ⚙️ TELA 4: CUSTOS FIXOS E RECÁLCULO INTELIGENTE
# =============================================================================
elif opcao_menu == "⚙️ 4. Custos Fixos & BD":
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("#### 💰 Hora-Máquina (R$/h)")
        df_maq = pd.DataFrame(list(valores_maquinas.items()), columns=["Máquina", "Taxa (R$/h)"])
        df_maq_edit = st.data_editor(
            df_maq, hide_index=True, use_container_width=True,
            column_config={"Máquina": st.column_config.TextColumn(disabled=True), "Taxa (R$/h)": st.column_config.NumberColumn(min_value=0.0, step=5.0, format="R$ %.2f")}
        )
        
    with c2:
        st.markdown("#### 📝 Impostos (%)")
        df_imp = pd.DataFrame(list(impostos.items()), columns=["Imposto", "Taxa (%)"])
        df_imp_edit = st.data_editor(
            df_imp, hide_index=True, use_container_width=True,
            column_config={"Imposto": st.column_config.TextColumn(disabled=True), "Taxa (%)": st.column_config.NumberColumn(min_value=0.0, step=0.5)}
        )

    if st.button("💾 Salvar Taxas e Impostos", type="primary"):
        novas_maq = dict(zip(df_maq_edit["Máquina"], df_maq_edit["Taxa (R$/h)"]))
        novos_imp = dict(zip(df_imp_edit["Imposto"], df_imp_edit["Taxa (%)"]))
        st.session_state.valores_maquinas = novas_maq
        st.session_state.impostos = novos_imp
        
        rows_mq = []
        for k, v in novas_maq.items(): rows_mq.append({"Tipo": "Máquina", "Nome": k, "Valor": float(v)})
        for k, v in novos_imp.items(): rows_mq.append({"Tipo": "Imposto", "Nome": k, "Valor": float(v)})
        salvar_aba_sheets(pd.DataFrame(rows_mq), "Config_Maquinas")
        st.toast("✅ Tarifas salvas com sucesso em Nuvem!")
        st.rerun()

    st.markdown("---")
    st.subheader("🔄 Recálculo de Tarifas (Hora-Máquina e Impostos)")
    if not df_init.empty:
        opcoes_motivo = [""] + ["Recálculo Geral (Use para Impostos ou Múltiplas Máquinas)"] + [f"Máquina: {m}" for m in valores_maquinas.keys()]
        motivo_recalc = st.selectbox("Qual tarifa motivou este recálculo?", opcoes_motivo)
        
        if motivo_recalc:
            clientes_selecionados = st.multiselect("Filtrar clientes (vazio = TODOS):", options=lista_empresas)
            df_ultimos = df_init.copy()
            df_ultimos['_dt_temp'] = pd.to_datetime(df_ultimos["Data/Hora"], format="%d/%m/%Y %H:%M", errors='coerce')
            df_ultimos = df_ultimos.sort_values("_dt_temp").groupby("Código da Peça").last().reset_index()
            
            if clientes_selecionados: 
                df_ultimos = df_ultimos[df_ultimos["Empresa"].isin(clientes_selecionados)]
                
            if "Geral" not in motivo_recalc:
                maq_alvo = motivo_recalc.replace("Máquina: ", "")
                df_ultimos = df_ultimos[df_ultimos["Usinagem_JSON"].apply(lambda x: maq_alvo in str(x))]
            
            if df_ultimos.empty: st.info("Nenhuma peça encontrada com estes critérios.")
            else:
                st.warning("⚠️ O recálculo utilizará as taxas VIVAS inseridas acima nas caixas.")
                df_ultimos.insert(0, "Recalcular", True)
                
                df_recalc_edit = st.data_editor(
                    df_ultimos[["Recalcular", "Empresa", "Código da Peça", "Nome da Peça", "Lote"]],
                    hide_index=True, use_container_width=True,
                    column_config={
                        "Recalcular": st.column_config.CheckboxColumn("Recalcular?"),
                        "Empresa": st.column_config.TextColumn(disabled=True),
                        "Código da Peça": st.column_config.TextColumn(disabled=True),
                        "Nome da Peça": st.column_config.TextColumn(disabled=True),
                        "Lote": st.column_config.NumberColumn(disabled=True)
                    }
                )
                
                if st.button("🚀 Executar Recálculo de Tarifas"):
                    # A CEREJA DO BOLO: Puxa os dados digitados "vivos" na tela imediatamente
                    novas_maq = dict(zip(df_maq_edit["Máquina"], df_maq_edit["Taxa (R$/h)"]))
                    novos_imp = dict(zip(df_imp_edit["Imposto"], df_imp_edit["Taxa (%)"]))
                    st.session_state.valores_maquinas = novas_maq
                    st.session_state.impostos = novos_imp
                    
                    rows_mq = []
                    for k, v in novas_maq.items(): rows_mq.append({"Tipo": "Máquina", "Nome": k, "Valor": float(v)})
                    for k, v in novos_imp.items(): rows_mq.append({"Tipo": "Imposto", "Nome": k, "Valor": float(v)})
                    salvar_aba_sheets(pd.DataFrame(rows_mq), "Config_Maquinas")

                    codigos_selecionados = df_recalc_edit[df_recalc_edit["Recalcular"]]["Código da Peça"].tolist()
                    linhas_recalculadas = []
                    
                    for _, row in df_ultimos[df_ultimos["Código da Peça"].isin(codigos_selecionados)].iterrows():
                        try:
                            row_clean = row.to_dict()
                            roteiro_raw = row_clean.get("Usinagem_JSON", "[]")
                            roteiro = json.loads(roteiro_raw) if isinstance(roteiro_raw, str) else []
                            
                            novo_custo_usinagem = 0.0
                            lote_recalc = max(1, safe_int(row_clean.get("Lote", 100)))
                            
                            if isinstance(roteiro, list):
                                for op in roteiro:
                                    pcs_h = max(0.1, safe_float(op.get("Peças por Hora"), 50.0))
                                    maq_nome = safe_str(op.get("Máquina"), "Outro")
                                    novo_custo_usinagem += (lote_recalc / pcs_h) * novas_maq.get(maq_nome, 120.0)
                                    op["Preço/Hora_Aplicado"] = novas_maq.get(maq_nome, 120.0)
                                row_clean["Usinagem_JSON"] = json.dumps(roteiro)
                            
                            custo_fabrica = safe_float(row_clean.get("Custo Material (R$)")) + novo_custo_usinagem + safe_float(row_clean.get("Custo Tratamento (R$)"))
                            fator_m = max(0.05, (1 - (safe_float(row_clean.get("Margem Lucro (%)", 30)) / 100)))
                            valor_venda = custo_fabrica / fator_m
                            imp_pct = safe_float(novos_imp.get("IR")) + safe_float(novos_imp.get("FS"))
                            preco_lote = valor_venda + (valor_venda * (imp_pct / 100))
                            
                            row_clean["Preço Total Lote (R$)"] = round(preco_lote, 2)
                            row_clean["Preço Unitário (R$)"] = round(preco_lote / lote_recalc, 2)
                            row_clean["Origem/Alteração"] = "Recálculo de Tarifas"
                            row_clean["Data/Hora"] = datetime.now().strftime("%d/%m/%Y %H:%M")
                            
                            if "Recalcular" in row_clean: del row_clean["Recalcular"]
                            if "_dt_temp" in row_clean: del row_clean["_dt_temp"]
                            linhas_recalculadas.append(row_clean)
                        except: pass
                        
                    if linhas_recalculadas:
                        df_final_salvar = pd.concat([df_init, pd.DataFrame(linhas_recalculadas)], ignore_index=True)
                        sucesso_tar = salvar_aba_sheets(df_final_salvar, "Historico")
                        if sucesso_tar:
                            st.session_state["db_historico"] = df_final_salvar
                            st.session_state["msg_sucesso_aba4"] = f"✅ {len(linhas_recalculadas)} orçamentos recalculados e consolidados em Nuvem!"
                            st.rerun()

    if st.session_state.get("msg_sucesso_aba4"):
        st.success(st.session_state["msg_sucesso_aba4"])
        st.session_state["msg_sucesso_aba4"] = ""

    st.markdown("---")
    st.subheader("🗑️ Zona de Perigo")
    if not df_init.empty:
        codigo_del = st.selectbox("Apagar histórico de um código específico:", [""] + lista_codigos)
        if codigo_del and st.button(f"Deletar todos os registros de {codigo_del}", type="primary"):
            df_restante = df_init[df_init["Código da Peça"] != codigo_del]
            sucesso_del = salvar_aba_sheets(df_restante, "Historico")
            if sucesso_del:
                st.session_state["db_historico"] = df_restante
                st.rerun()
            
# =============================================================================
# 🛒 TELA 5: COTAÇÃO FINAL (VENDAS)
# =============================================================================
elif opcao_menu == "🛒 5. Cotação Final":
    st.subheader("🛒 Módulo de Vendas (Simulador de Cotação)")
    if df_init.empty:
        st.info("Nenhum orçamento disponível no banco de dados da engenharia.")
    else:
        st.markdown("### Passo 1: Seleção do Cliente e Peças")
        cliente_cot = st.selectbox("Selecione o Cliente para cotar:", [""] + lista_empresas)

        if cliente_cot:
            df_cli = df_init[df_init["Empresa"] == cliente_cot].copy()
            df_cli['_dt'] = pd.to_datetime(df_cli["Data/Hora"], format="%d/%m/%Y %H:%M", errors='coerce')
            df_ultimos = df_cli.sort_values("_dt").groupby("Código da Peça").last().reset_index()

            pecas_selecionadas = st.multiselect("Adicione as peças ao carrinho:", options=df_ultimos["Código da Peça"].tolist())

            if pecas_selecionadas:
                st.markdown("---")
                st.markdown("### Passo 2: Mesa de Negociação")
                dados_carrinho = []
                for _, row in df_ultimos[df_ultimos["Código da Peça"].isin(pecas_selecionadas)].iterrows():
                    r = row.to_dict()
                    lote_original = max(1, safe_int(r.get("Lote", 100)))
                    custo_mat_unit = safe_float(r.get("Custo Material (R$)")) / lote_original
                    custo_trat_unit = safe_float(r.get("Custo Tratamento (R$)")) / lote_original
                    peso_unit = safe_float(r.get("Total Kg Lote")) / lote_original

                    roteiro_raw = r.get("Usinagem_JSON", "[]")
                    roteiro = json.loads(roteiro_raw) if isinstance(roteiro_raw, str) else []
                    custo_usi_unit = 0.0
                    if isinstance(roteiro, list):
                        for op in roteiro:
                            pcs_h = max(0.1, safe_float(op.get("Peças por Hora"), 50.0))
                            maq_nome = safe_str(op.get("Máquina"), "Outro")
                            taxa_hora = safe_float(op.get("Preço/Hora_Aplicado", valores_maquinas.get(maq_nome, 120.0)))
                            custo_usi_unit += (1 / pcs_h) * taxa_hora

                    custo_base_unit = custo_mat_unit + custo_trat_unit + custo_usi_unit

                    dados_carrinho.append({
                        "Código da Peça": r["Código da Peça"], "Nome": safe_str(r.get("Nome da Peça", "")),
                        "Custo Base (R$)": custo_base_unit, "Peso Unit. (Kg)": peso_unit,
                        "Novo Lote": lote_original, "Margem Desejada (%)": safe_float(r.get("Margem Lucro (%)", 30))
                    })

                df_base_cotacao = pd.DataFrame(dados_carrinho)
                df_editado = st.data_editor(
                    df_base_cotacao, hide_index=True, use_container_width=True,
                    column_config={
                        "Código da Peça": st.column_config.TextColumn(disabled=True), "Nome": st.column_config.TextColumn(disabled=True),
                        "Custo Base (R$)": st.column_config.NumberColumn("Custo Base (Não Editável)", format="R$ %.2f", disabled=True),
                        "Peso Unit. (Kg)": st.column_config.NumberColumn(format="%.3f kg", disabled=True),
                        "Novo Lote": st.column_config.NumberColumn(min_value=1, step=1),
                        "Margem Desejada (%)": st.column_config.NumberColumn(min_value=1.0, max_value=99.0, step=1.0)
                    }
                )

                total_imposto_pct = safe_float(st.session_state.impostos.get("IR")) + safe_float(st.session_state.impostos.get("FS"))
                valor_total_cotacao, imposto_total_cotacao, custo_total_cotacao, peso_total_cotacao = 0.0, 0.0, 0.0, 0.0
                resultados_visuais = []

                for _, row in df_editado.iterrows():
                    lote = safe_int(row["Novo Lote"])
                    margem = safe_float(row["Margem Desejada (%)"])
                    custo_unit = safe_float(row["Custo Base (R$)"])
                    peso_total_cotacao += safe_float(row["Peso Unit. (Kg)"]) * lote

                    custo_bruto_lote = custo_unit * lote
                    fator_lucro = max(0.01, (1 - (margem / 100)))
                    valor_sem_imposto = custo_bruto_lote / fator_lucro
                    imposto_reais = valor_sem_imposto * (total_imposto_pct / 100)
                    preco_lote_final = valor_sem_imposto + imposto_reais
                    preco_unit_final = preco_lote_final / lote

                    valor_total_cotacao += preco_lote_final
                    imposto_total_cotacao += imposto_reais
                    custo_total_cotacao += custo_bruto_lote

                    resultados_visuais.append({
                        "Código": row["Código da Peça"], "Nome": row["Nome"], "Lote": lote,
                        "Preço de Venda Unitário": round(preco_unit_final, 2), "Total Faturado": round(preco_lote_final, 2)
                    })

                st.markdown("---")
                st.markdown("### Passo 3: A Proposta (Visão do Cliente)")
                
                def formatar_br(valor): return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                
                df_resumo = pd.DataFrame(resultados_visuais)
                # AJUSTE DA VÍRGULA: Formatação regional comercial aplicada na visualização
                st.dataframe(
                    df_resumo, hide_index=True, use_container_width=True,
                    column_config={
                        "Preço de Venda Unitário": st.column_config.NumberColumn("Preço Unitário", format="R$ %.2f"),
                        "Total Faturado": st.column_config.NumberColumn("Total Faturado", format="R$ %.2f")
                    }
                )
                
                st.markdown("---")
                st.markdown("### 📊 Passo 4: Raio-X da Negociação")
                lucro_bruto_reais = valor_total_cotacao - custo_total_cotacao
                lucro_real_reais = lucro_bruto_reais - imposto_total_cotacao
                
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("1. Faturamento", f"R$ {formatar_br(valor_total_cotacao)}")
                c2.metric("2. Custo de Fábrica", f"R$ {formatar_br(custo_total_cotacao)}")
                c3.metric(f"3. Impostos ({total_imposto_pct}%)", f"R$ {formatar_br(imposto_total_cotacao)}")
                c4.metric("4. Lucro Real", f"R$ {formatar_br(lucro_real_reais)}")

                st.markdown("---")
                st.markdown("### Passo 5: Salvar e Exportar")
                
                agora = datetime.now()
                ano_atual = agora.year
                data_str = agora.strftime("%d/%m/%Y")
                df_cot_nuvem = st.session_state["db_cotacoes"]
                
                novo_id = f"COT-{ano_atual}-001"
                if not df_cot_nuvem.empty and "ID Cotação" in df_cot_nuvem.columns:
                    ids_ano = df_cot_nuvem[df_cot_nuvem['ID Cotação'].astype(str).str.contains(f"COT-{ano_atual}", na=False)]
                    if not ids_ano.empty:
                        max_num = ids_ano['ID Cotação'].str.split('-').str[-1].astype(int).max()
                        novo_id = f"COT-{ano_atual}-{max_num + 1:03d}"

                def gerar_doc_word(id_cot, cliente, df_itens, total, peso):
                    doc = Document()
                    doc.add_heading('LENOOR ME - Proposta Comercial', 0)
                    doc.add_paragraph(f"Número da Proposta: {id_cot}\nData: {data_str}\n\nÀ Empresa: {cliente}")
                    table = doc.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text, hdr_cells[4].text = 'Código', 'Descrição', 'Qtd', 'Preço Unit.', 'Total'
                    
                    for _, row in df_itens.iterrows():
                        row_cells = table.add_row().cells
                        row_cells[0].text, row_cells[1].text, row_cells[2].text = str(row['Código']), str(row['Nome']), str(row['Lote'])
                        row_cells[3].text = f"R$ {formatar_br(row['Preço de Venda Unitário'])}"
                        row_cells[4].text = f"R$ {formatar_br(row['Total Faturado'])}"
                        
                    doc.add_paragraph(f"\nValor Total da Proposta: R$ {formatar_br(total)}")
                    doc.add_paragraph(f"Peso Estimado da Carga: {formatar_br(peso)} kg\n")
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    return buffer

                col_b1, col_b2 = st.columns(2)
                with col_b1:
                    if st.button(f"💾 Salvar Proposta Comercial ({novo_id})", type="primary", use_container_width=True):
                        df_salvar_cot = df_resumo.copy()
                        df_salvar_cot.insert(0, "ID Cotação", novo_id)
                        df_salvar_cot.insert(1, "Data", data_str)
                        df_salvar_cot.insert(2, "Cliente", cliente_cot)
                        df_salvar_cot["Valor Total Proposta"] = valor_total_cotacao
                        
                        df_final_cot = pd.concat([df_cot_nuvem, df_salvar_cot], ignore_index=True) if not df_cot_nuvem.empty else df_salvar_cot
                        sucesso_cot = salvar_aba_sheets(df_final_cot, "Historico_Cotacoes")
                        if sucesso_cot:
                            st.session_state["db_cotacoes"] = df_final_cot
                            st.success(f"✅ {novo_id} gravada permanentemente na Nuvem!")
                        
                with col_b2:
                    arquivo_word = gerar_doc_word(novo_id, cliente_cot, df_resumo, valor_total_cotacao, peso_total_cotacao)
                    st.download_button(label="📄 Baixar Proposta Padrão (.docx)", data=arquivo_word, file_name=f"{novo_id}_{cliente_cot}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

# =============================================================================
# 📁 TELA 6: HISTÓRICO DE COTAÇÕES (CRM IN CLOUD)
# =============================================================================
elif opcao_menu == "📁 6. Histórico de Cotações":
    st.subheader("📁 Central de Propostas Comerciais")
    df_historico = st.session_state["db_cotacoes"]
    
    if df_historico.empty:
        st.info("Nenhuma cotação salva na planilha em nuvem ainda.")
    else:
        st.markdown("#### 🔍 Filtros de Busca")
        c_filt1, c_filt2 = st.columns(2)
        with c_filt1:
            clientes_cot = ["Todos"] + sorted(df_historico["Cliente"].dropna().astype(str).unique().tolist())
            filtro_cliente = st.selectbox("Filtrar por Cliente:", clientes_cot)
        with c_filt2:
            ids_cot = ["Todas"] + sorted(df_historico["ID Cotação"].dropna().astype(str).unique().tolist(), reverse=True)
            filtro_id = st.selectbox("Filtrar por Número (ID):", ids_cot)
            
        df_filtrado = df_historico.copy()
        if filtro_cliente != "Todos": df_filtrado = df_filtrado[df_filtrado["Cliente"] == filtro_cliente]
        if filtro_id != "Todas": df_filtrado = df_filtrado[df_filtrado["ID Cotação"] == filtro_id]

        st.markdown("---")
        st.markdown("#### 📑 Visão Geral das Propostas")
        df_master = df_filtrado.groupby(["ID Cotação", "Data", "Cliente"]).agg({"Valor Total Proposta": "max"}).reset_index().sort_values("ID Cotação", ascending=False)
        
        # AJUSTE DA VÍRGULA: Formatação visual de moeda perfeita aplicada no CRM
        st.dataframe(
            df_master, hide_index=True, use_container_width=True,
            column_config={"Valor Total Proposta": st.column_config.NumberColumn("Valor Total", format="R$ %.2f")}
        )
        
        st.markdown("---")
        st.markdown("#### 🔍 Detalhes da Proposta Selecionada")
        id_selecionado = st.selectbox("Selecione a Cotação para ver os itens:", df_master["ID Cotação"].tolist())
        
        if id_selecionado:
            df_detalhe = df_historico[df_historico["ID Cotação"] == id_selecionado][["Código", "Nome", "Lote", "Preço de Venda Unitário", "Total Faturado"]]
            
            st.dataframe(
                df_detalhe, hide_index=True, use_container_width=True,
                column_config={
                    "Preço de Venda Unitário": st.column_config.NumberColumn("Preço Unitário", format="R$ %.2f"),
                    "Total Faturado": st.column_config.NumberColumn("Total Faturado", format="R$ %.2f")
                }
            )
            
            buffer_t6 = io.BytesIO()
            with pd.ExcelWriter(buffer_t6, engine='xlsxwriter') as writer:
                df_detalhe.to_excel(writer, sheet_name=id_selecionado, index=False)
            st.download_button("📥 Exportar Itens para Excel (.xlsx)", data=buffer_t6.getvalue(), file_name=f"{id_selecionado}_itens.xlsx", mime="application/vnd.ms-excel")
